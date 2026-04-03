"""
SRS File Reader — extracts plain text from PDF, DOCX, or TXT files.
Libraries: pdfplumber (PDF), python-docx (DOCX)
"""
import io
from typing import Union


def read_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber (handles tables + text)."""
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            text_parts.append(f"\n--- Page {page_num} ---")

            # Extract tables first (structured format so Gemini identifies headers)
            tables = page.extract_tables()
            for table in tables:
                if not table:
                    continue
                header = table[0]
                text_parts.append("TABLE: " + " | ".join(cell or "" for cell in header))
                for row in table[1:]:
                    text_parts.append(" | ".join(cell or "" for cell in row))

            # Extract plain text
            text = page.extract_text()
            if text:
                text_parts.append(text)

    return "\n".join(text_parts)


def read_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx (handles paragraphs + tables)."""
    from docx import Document

    text_parts = []
    doc = Document(io.BytesIO(file_bytes))

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                text_parts.append(row_text)

    return "\n".join(text_parts)


def read_txt(file_bytes: bytes) -> str:
    """Read plain text file."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")


def extract_text(filename: str, file_bytes: bytes) -> str:
    """
    Auto-detect file type by extension and extract text.
    Returns clean text string.
    """
    ext = filename.lower().rsplit(".", 1)[-1]

    if ext == "pdf":
        return read_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return read_docx(file_bytes)
    elif ext == "txt":
        return read_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext} — supported: pdf, docx, txt")
