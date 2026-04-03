"""
CAF SRS → Professional DOCX Generator
Generates a techno-functional SRS document from the structured TXT format.
Output: CAF_Full_SRS_v2.docx (suitable for print / PDF export)

Usage:
    python generate_srs_docx.py
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Color Palette ────────────────────────────────────────────────────────────
COLOR_PRIMARY      = RGBColor(0xE9, 0x09, 0x0C)   # Karnataka red
COLOR_DARK         = RGBColor(0x1A, 0x1A, 0x2E)   # near-black heading
COLOR_SECTION_BG   = RGBColor(0xF0, 0xF4, 0xFA)   # light blue-grey
COLOR_PAGE_BG      = RGBColor(0xE9, 0x09, 0x0C)   # red page header
COLOR_FIELD_EVEN   = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_FIELD_ODD    = RGBColor(0xFD, 0xFD, 0xFD)
COLOR_MANDATORY_Y  = RGBColor(0xD3, 0x29, 0x29)   # red
COLOR_MANDATORY_N  = RGBColor(0x6C, 0x75, 0x7D)   # grey
COLOR_MANDATORY_C  = RGBColor(0xE6, 0x72, 0x00)   # orange
COLOR_TYPE_BG      = {
    "text":       RGBColor(0xE8, 0xF5, 0xE9),
    "select":     RGBColor(0xE3, 0xF2, 0xFD),
    "radio":      RGBColor(0xF3, 0xE5, 0xF5),
    "checkbox":   RGBColor(0xFD, 0xF6, 0xFF),
    "date":       RGBColor(0xFD, 0xF3, 0xE7),
    "tel":        RGBColor(0xE0, 0xF7, 0xFA),
    "email":      RGBColor(0xE0, 0xF7, 0xFA),
    "textarea":   RGBColor(0xF9, 0xFB, 0xE7),
    "file":       RGBColor(0xFF, 0xF8, 0xE1),
    "number":     RGBColor(0xE8, 0xF5, 0xE9),
    "multiselect":RGBColor(0xE3, 0xF2, 0xFD),
}

TYPE_DISPLAY = {
    "text": "Text",        "select": "Dropdown",    "radio": "Radio",
    "checkbox": "Checkbox","date": "Date",          "tel": "Mobile",
    "email": "Email",      "textarea": "Textarea",  "file": "File Upload",
    "number": "Number",    "multiselect": "Multi-select",
}

MANDATORY_DISPLAY = {
    "Yes": ("Mandatory", COLOR_MANDATORY_Y),
    "No":  ("Optional",  COLOR_MANDATORY_N),
    "Conditional": ("Conditional", COLOR_MANDATORY_C),
    "Auto Populated": ("Auto Populated", COLOR_MANDATORY_N),
    "Auto Calculated": ("Auto Calculated", COLOR_MANDATORY_N),
}


# ─── XML helpers ──────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, spec in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if spec:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   spec.get('val', 'single'))
            el.set(qn('w:sz'),    str(spec.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), spec.get('color', 'auto'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_run_colored(para, text, bold=False, italic=False,
                    font_size=10, color: RGBColor = None, font_name="Calibri"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    return run


def no_space_before_after(para):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'),  '0')
    pPr.append(spacing)


# ─── Document setup ───────────────────────────────────────────────────────────

def setup_document() -> Document:
    doc = Document()
    # Narrow margins
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)
    # Remove default paragraph spacing
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    return doc


# ─── Cover Page ───────────────────────────────────────────────────────────────

def add_cover_page(doc: Document):
    # Red top banner
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COLOR_PAGE_BG)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(14)
    r = p.add_run("GOVERNMENT OF KARNATAKA — DEPARTMENT OF INDUSTRIES & COMMERCE")
    r.font.name  = "Calibri"
    r.font.size  = Pt(11)
    r.bold       = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_space_before_after(p)
    r = p.add_run("Combined Application Form (CAF)")
    r.font.name  = "Calibri"
    r.font.size  = Pt(26)
    r.bold       = True
    r.font.color.rgb = COLOR_PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_space_before_after(p)
    r = p.add_run("Service Requirement Specification (SRS)")
    r.font.name  = "Calibri"
    r.font.size  = Pt(16)
    r.bold       = True
    r.font.color.rgb = COLOR_DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Techno-Functional Document — FormBuilder System")
    r.font.size  = Pt(11)
    r.italic     = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # Metadata table
    meta_data = [
        ("Form Name",        "Combined Application Form (CAF)"),
        ("Service ID",       "591.0"),
        ("Department ID",    "1"),
        ("Form Type ID",     "1  (New Application)"),
        ("Version",          "2.0 — Final"),
        ("Date",             "2026-03-12"),
        ("Total Pages",      "7 (+ 3 system pages: Payment, Signing, Summary)"),
        ("Total Fields",     "143 fields | 25 conditional rules | 1 add-more group"),
        ("Purpose",          "Establishing new / Expansion / Diversification /\n"
                             "Modernization of Industrial Units in Karnataka"),
    ]
    tbl = doc.add_table(rows=len(meta_data), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    for i, (key, val) in enumerate(meta_data):
        r0 = tbl.cell(i, 0)
        r1 = tbl.cell(i, 1)
        set_cell_bg(r0, RGBColor(0xF5, 0xF5, 0xF5))
        r0.width = Inches(2.2)
        r1.width = Inches(4.3)
        p0 = r0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(3)
        p0.paragraph_format.space_after  = Pt(3)
        add_run_colored(p0, key, bold=True, font_size=10, color=COLOR_DARK)
        p1 = r1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(3)
        p1.paragraph_format.space_after  = Pt(3)
        add_run_colored(p1, val, font_size=10)

    doc.add_paragraph()

    # Subtitle note
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("This document is the authoritative SRS for AI-assisted FormBuilder JSON generation.\n"
                  "Follow this exact format when onboarding new department services.")
    r.font.size  = Pt(9)
    r.italic     = True
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_page_break()


# ─── How To Use Section ───────────────────────────────────────────────────────

def add_how_to_use(doc: Document):
    add_heading1(doc, "HOW TO USE THIS SRS DOCUMENT")

    p = doc.add_paragraph()
    r = p.add_run(
        "This SRS (Service Requirement Specification) document is the single source of truth "
        "for configuring any government service form in the Karnataka Single Window FormBuilder system. "
        "When uploaded to the AI FormBuilder, it auto-generates a complete database-ready JSON."
    )
    r.font.size = Pt(10)

    add_heading2(doc, "1. Purpose of This Document")
    bullets = [
        "Define every field, type, validation, and conditional rule for the CAF form",
        "Serve as input to AI FormBuilder to generate DB-ready JSON (no manual SQL needed)",
        "Act as a template for onboarding new department services",
        "Provide a techno-functional reference for developers, BAs, and department officials",
    ]
    for b in bullets:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        add_run_colored(p, b, font_size=10)

    add_heading2(doc, "2. How AI FormBuilder Uses This SRS")
    steps = [
        ("Step 1", "Upload this .txt file to the AI FormBuilder system"),
        ("Step 2", "AI reads every PAGE, SECTION, FIELD, and CONDITIONAL RULE"),
        ("Step 3", "AI generates a single structured JSON matching the DB schema"),
        ("Step 4", "Admin reviews JSON in preview screen"),
        ("Step 5", "On approval, JSON is inserted into 10 m_fb_* database tables"),
        ("Step 6", "Form is live for investors on the Single Window Portal"),
    ]
    for num, desc in steps:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        add_run_colored(p, f"  {num}: ", bold=True, font_size=10, color=COLOR_PRIMARY)
        add_run_colored(p, desc, font_size=10)

    add_heading2(doc, "3. Field Type Reference")
    type_data = [
        ("text",        "Single-line text input",                    "Name, Address, PAN, CIN"),
        ("number",      "Numeric input (integer or decimal)",         "Investment, Area, Employment"),
        ("select",      "Single-select dropdown",                     "District, Type Of Proposal"),
        ("multiselect", "Multi-value dropdown",                       "Assistance Type"),
        ("radio",       "Radio button (2-4 options)",                 "Yes/No, New/Existing"),
        ("checkbox",    "Single checkbox (true/false)",               "Same Address flag"),
        ("date",        "Date picker (calendar)",                     "Date Of Incorporation"),
        ("tel",         "Phone/mobile number input",                  "Mobile Number"),
        ("email",       "Email address input",                        "Corporate Email"),
        ("textarea",    "Multi-line text area",                       "Products Description"),
        ("file",        "File upload (DMS integrated)",               "PAN Card, DPR, Land Deed"),
        ("hidden",      "Hidden system field (not shown to user)",    "Internal tracking IDs"),
    ]
    tbl = doc.add_table(rows=1 + len(type_data), cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Field Type", "Description", "Examples in CAF"]
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        set_cell_bg(cell, COLOR_DARK)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        add_run_colored(p, h, bold=True, font_size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
    for row_i, (ftype, fdesc, fex) in enumerate(type_data):
        bg = TYPE_DISPLAY.get(ftype)
        row_bg = COLOR_FIELD_ODD if row_i % 2 == 0 else COLOR_FIELD_EVEN
        for col_i, val in enumerate([ftype.upper(), fdesc, fex]):
            cell = tbl.cell(row_i + 1, col_i)
            set_cell_bg(cell, row_bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            is_bold = (col_i == 0)
            col_color = COLOR_PRIMARY if col_i == 0 else COLOR_DARK
            add_run_colored(p, val, bold=is_bold, font_size=9, color=col_color)

    add_heading2(doc, "4. Mandatory Field Values")
    mand_info = [
        ("Yes",           COLOR_MANDATORY_Y, "Always required. Form cannot be submitted without this field."),
        ("No",            COLOR_MANDATORY_N, "Optional. User may leave blank."),
        ("Conditional",   COLOR_MANDATORY_C, "Required only when specific condition is met (see CONDITIONAL RULES)."),
        ("Auto Populated",COLOR_MANDATORY_N, "Pre-filled from investor registration. Usually read-only."),
        ("Auto Calculated",COLOR_MANDATORY_N,"System calculates value (e.g. totals). Always read-only."),
    ]
    for mval, mcol, mdesc in mand_info:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        add_run_colored(p, f"  [{mval}]  ", bold=True, font_size=10, color=mcol)
        add_run_colored(p, mdesc, font_size=10)

    add_heading2(doc, "5. Source Types for Dropdown/Select Fields")
    src_info = [
        ("STATIC",  "Options listed inline in this SRS. Fixed values (e.g. Yes/No, Zone 1-5)."),
        ("MASTER",  "Options come from a DB master table (e.g. Country, District, NIC Code).\n"
                    "           Master table ID is configured in FormBuilder admin."),
    ]
    for stype, sdesc in src_info:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        add_run_colored(p, f"  {stype}:  ", bold=True, font_size=10, color=COLOR_PRIMARY)
        add_run_colored(p, sdesc, font_size=10)

    add_heading2(doc, "6. How to Write SRS for a New Service")
    steps2 = [
        "Copy this document as a template",
        "Replace FORM METADATA (Form Name, Department ID, Service ID, Form Type ID)",
        "Define each PAGE with sections and fields in the same format",
        "For each field: Name | Type | Mandatory | Grid | Options/Source | Conditions | Validation | Help Text",
        "State CONDITIONAL RULES at the end of each section",
        "For repeatable rows: mark section as ADD MORE GROUP with min/max rows",
        "Upload the .txt version to AI FormBuilder to generate JSON automatically",
    ]
    for i, s in enumerate(steps2, 1):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        add_run_colored(p, s, font_size=10)

    doc.add_page_break()


# ─── Section Headings ─────────────────────────────────────────────────────────

def add_heading1(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(6)
    # red left border via shading the whole para — use a 1-cell table trick
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COLOR_PAGE_BG)
    set_cell_borders(cell,
        top    = {'val': 'none'},
        bottom = {'val': 'none'},
        left   = {'val': 'none'},
        right  = {'val': 'none'},
    )
    ph = cell.paragraphs[0]
    ph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ph.paragraph_format.space_before = Pt(6)
    ph.paragraph_format.space_after  = Pt(6)
    r = ph.add_run(f"  {text}")
    r.font.name  = "Calibri"
    r.font.size  = Pt(13)
    r.bold       = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return tbl


def add_heading2(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(11)
    r.bold       = True
    r.font.color.rgb = COLOR_DARK
    return p


def add_page_header(doc: Document, page_num: int, page_title: str):
    """Red banner for each PAGE section."""
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Left: PAGE label
    c0 = tbl.cell(0, 0)
    c1 = tbl.cell(0, 1)
    set_cell_bg(c0, COLOR_PAGE_BG)
    set_cell_bg(c1, COLOR_PAGE_BG)
    c0.width = Inches(1.0)
    c1.width = Inches(5.5)

    p0 = c0.paragraphs[0]
    p0.paragraph_format.space_before = Pt(5)
    p0.paragraph_format.space_after  = Pt(5)
    r = p0.add_run(f"  PAGE {page_num}")
    r.font.name  = "Calibri"
    r.font.size  = Pt(11)
    r.bold       = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p1 = c1.paragraphs[0]
    p1.paragraph_format.space_before = Pt(5)
    p1.paragraph_format.space_after  = Pt(5)
    r1 = p1.add_run(f"  {page_title.upper()}")
    r1.font.name  = "Calibri"
    r1.font.size  = Pt(11)
    r1.bold       = True
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def add_section_header(doc: Document, section_label: str, section_title: str, note: str = None):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COLOR_SECTION_BG)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    add_run_colored(p, f"  SECTION {section_label} — ", bold=True, font_size=10,
                    color=COLOR_PRIMARY)
    add_run_colored(p, section_title, bold=True, font_size=10, color=COLOR_DARK)
    if note:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after  = Pt(4)
        add_run_colored(p2, f"  ℹ  {note}", italic=True, font_size=9,
                        color=RGBColor(0x55, 0x66, 0x88))


# ─── Field Table ──────────────────────────────────────────────────────────────

FIELD_TABLE_HEADERS = ["#", "Field Name", "Type", "Grid", "Mandatory", "Validation / Notes"]
FIELD_COL_WIDTHS    = [Inches(0.35), Inches(1.8), Inches(0.85), Inches(0.4), Inches(0.85), Inches(2.7)]

def start_field_table(doc: Document):
    tbl = doc.add_table(rows=1, cols=len(FIELD_TABLE_HEADERS))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (h, w) in enumerate(zip(FIELD_TABLE_HEADERS, FIELD_COL_WIDTHS)):
        cell = tbl.cell(0, i)
        cell.width = w
        set_cell_bg(cell, COLOR_DARK)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        add_run_colored(p, h, bold=True, font_size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
    return tbl


def add_field_row(tbl, field_num: str, name: str, ftype: str, grid: str,
                  mandatory: str, notes: str, row_idx: int):
    row = tbl.add_row()
    bg = COLOR_FIELD_ODD if row_idx % 2 == 0 else COLOR_FIELD_EVEN
    type_bg = TYPE_DISPLAY.get(ftype.lower(), ftype)

    # Col 0: Number
    c0 = row.cells[0]
    c0.width = FIELD_COL_WIDTHS[0]
    set_cell_bg(c0, bg)
    p = c0.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    add_run_colored(p, field_num, font_size=8, color=RGBColor(0x88, 0x88, 0x88))

    # Col 1: Name
    c1 = row.cells[1]
    c1.width = FIELD_COL_WIDTHS[1]
    set_cell_bg(c1, bg)
    p = c1.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    add_run_colored(p, name, bold=True, font_size=9, color=COLOR_DARK)

    # Col 2: Type
    c2 = row.cells[2]
    c2.width = FIELD_COL_WIDTHS[2]
    type_key = ftype.lower()
    type_cell_bg = TYPE_DISPLAY.get(type_key, "")
    type_color_map = {
        "text": RGBColor(0x2E, 0x7D, 0x32),    "select": RGBColor(0x15, 0x65, 0xC0),
        "radio": RGBColor(0x6A, 0x1B, 0x9A),   "checkbox": RGBColor(0x6A, 0x1B, 0x9A),
        "date": RGBColor(0xE6, 0x51, 0x00),    "tel": RGBColor(0x00, 0x6E, 0x7F),
        "email": RGBColor(0x00, 0x6E, 0x7F),   "textarea": RGBColor(0x55, 0x5D, 0x00),
        "file": RGBColor(0x99, 0x6A, 0x00),    "number": RGBColor(0x2E, 0x7D, 0x32),
        "multiselect": RGBColor(0x15, 0x65, 0xC0),
    }
    set_cell_bg(c2, bg)
    p = c2.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    disp_type = TYPE_DISPLAY.get(type_key, ftype.title())
    t_color   = type_color_map.get(type_key, COLOR_DARK)
    add_run_colored(p, disp_type, bold=True, font_size=9, color=t_color)

    # Col 3: Grid
    c3 = row.cells[3]
    c3.width = FIELD_COL_WIDTHS[3]
    set_cell_bg(c3, bg)
    p = c3.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    g_color = COLOR_DARK if grid == "6" else COLOR_PRIMARY
    add_run_colored(p, grid, bold=(grid == "12"), font_size=9, color=g_color)

    # Col 4: Mandatory
    c4 = row.cells[4]
    c4.width = FIELD_COL_WIDTHS[4]
    set_cell_bg(c4, bg)
    p = c4.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    mtext, mcolor = MANDATORY_DISPLAY.get(mandatory, (mandatory, COLOR_DARK))
    add_run_colored(p, mtext, bold=(mandatory == "Yes"), font_size=9, color=mcolor)

    # Col 5: Notes
    c5 = row.cells[5]
    c5.width = FIELD_COL_WIDTHS[5]
    set_cell_bg(c5, bg)
    p = c5.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    add_run_colored(p, notes, font_size=8.5, color=RGBColor(0x33, 0x33, 0x55))


def add_options_row(tbl, options_text: str, row_idx: int):
    """A full-width row showing static options or source."""
    row  = tbl.add_row()
    bg   = RGBColor(0xFD, 0xF8, 0xFF)
    # Merge all cols
    row.cells[0].merge(row.cells[len(FIELD_TABLE_HEADERS) - 1])
    cell = row.cells[0]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    add_run_colored(p, "   Options: ", bold=True, font_size=8.5, color=COLOR_PRIMARY)
    add_run_colored(p, options_text, font_size=8.5, color=RGBColor(0x44, 0x44, 0x66))


def add_condition_row(tbl, condition_text: str):
    """A full-width row showing conditional display rule."""
    row = tbl.add_row()
    bg  = RGBColor(0xFF, 0xFB, 0xEE)
    row.cells[0].merge(row.cells[len(FIELD_TABLE_HEADERS) - 1])
    cell = row.cells[0]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    add_run_colored(p, "   ⚡ Show when: ", bold=True, font_size=8.5,
                    color=RGBColor(0xBB, 0x6A, 0x00))
    add_run_colored(p, condition_text, font_size=8.5, color=RGBColor(0x55, 0x44, 0x00))


def add_addmore_banner(doc: Document, label: str, min_rows: int, max_rows: int):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, RGBColor(0x1A, 0x1A, 0x2E))
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    add_run_colored(p, f"  ➕  ADD MORE (Repeatable): {label}  "
                       f"  [Min rows: {min_rows}  |  Max rows: {max_rows}]",
                    bold=True, font_size=9.5, color=RGBColor(0xFF, 0xFF, 0xFF))


def add_rules_block(doc: Document, rules: list):
    if not rules:
        return
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, RGBColor(0xFF, 0xF9, 0xE8))
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(1)
    add_run_colored(p, "  CONDITIONAL RULES", bold=True, font_size=10,
                    color=RGBColor(0x99, 0x55, 0x00))
    for rule in rules:
        pr = doc.add_paragraph()
        pr.paragraph_format.space_before = Pt(2)
        pr.paragraph_format.space_after  = Pt(2)
        add_run_colored(pr, f"   {rule}", font_size=9, color=RGBColor(0x55, 0x44, 0x00))
    doc.add_paragraph()


# ─── SRS TXT Parser ───────────────────────────────────────────────────────────

class SRSParser:
    """Parse the structured TXT SRS into a Python data structure."""

    def __init__(self, txt: str):
        self.lines = txt.splitlines()
        self.pages = []
        self._parse()

    def _parse(self):
        pages = []
        current_page = None
        current_section = None
        current_field = None
        current_rules = []
        collecting_options = False
        options_lines = []
        collecting_condition = False
        condition_lines = []
        addmore_group = None

        def flush_field():
            nonlocal current_field, collecting_options, options_lines
            nonlocal collecting_condition, condition_lines
            if current_field and current_section is not None:
                if collecting_options and options_lines:
                    current_field['options'] = " | ".join(options_lines)
                if collecting_condition and condition_lines:
                    current_field['condition'] = " ".join(condition_lines)
                current_section['fields'].append(current_field)
            collecting_options = False
            options_lines = []
            collecting_condition = False
            condition_lines = []

        def flush_section():
            nonlocal current_section, current_rules
            if current_section and current_page is not None:
                flush_field()
                current_section['rules'] = current_rules[:]
                current_page['sections'].append(current_section)
            current_rules.clear()

        for raw in self.lines:
            line = raw.strip()

            # PAGE header
            m = re.match(r'^PAGE\s+(\d+)\s+[—–-]+\s*(.+)', line)
            if m:
                flush_section()
                if current_page:
                    pages.append(current_page)
                current_page    = {'num': int(m.group(1)), 'title': m.group(2).strip(),
                                   'sections': [], 'note': ''}
                current_section = None
                current_field   = None
                current_rules   = []
                addmore_group   = None
                continue

            if current_page is None:
                continue

            # SECTION header
            m = re.match(r'^SECTION\s+(\S+)\s+[—–-]+\s*(.+)', line)
            if m:
                flush_section()
                current_section = {'label': m.group(1), 'title': m.group(2).strip(),
                                   'fields': [], 'rules': [], 'note': '', 'addmore': None}
                current_field   = None
                current_rules   = []
                addmore_group   = None
                continue

            if current_section is None:
                if line.startswith('NOTE:'):
                    if current_page:
                        current_page['note'] += line[5:].strip() + ' '
                continue

            # Section NOTE
            if line.startswith('NOTE:'):
                current_section['note'] += line[5:].strip() + ' '
                continue

            # ADD MORE GROUP
            m = re.match(r'^ADD MORE GROUP:\s*(.+)', line)
            if m:
                addmore_group = {'label': m.group(1).strip(), 'min': 1, 'max': 10}
                current_section['addmore'] = addmore_group
                continue
            if addmore_group:
                m2 = re.match(r'Minimum rows:\s*(\d+)', line)
                if m2:
                    addmore_group['min'] = int(m2.group(1)); continue
                m2 = re.match(r'Maximum rows:\s*(\d+)', line)
                if m2:
                    addmore_group['max'] = int(m2.group(1)); continue

            # FIELD line
            m = re.match(r'^FIELD\s+(\d+)\s*\|?\s*Name:\s*(.+)', line)
            if m:
                flush_field()
                current_field = {
                    'num': m.group(1).strip(), 'name': m.group(2).strip(),
                    'type': '', 'mandatory': '', 'grid': '6', 'readonly': False,
                    'max_length': '', 'pattern': '', 'options': '', 'source': '',
                    'condition': '', 'validation': '', 'help_text': '', 'notes': '',
                }
                collecting_options   = False
                collecting_condition = False
                continue

            if current_field:
                # Type
                m = re.match(r'^Type:\s*(.+)', line)
                if m: current_field['type'] = m.group(1).strip().lower(); continue

                m = re.match(r'^Mandatory:\s*(.+)', line)
                if m: current_field['mandatory'] = m.group(1).strip(); continue

                m = re.match(r'^Grid:\s*(.+)', line)
                if m: current_field['grid'] = m.group(1).strip(); continue

                m = re.match(r'^Is Readonly:\s*(.+)', line)
                if m:
                    val = m.group(1).strip().lower()
                    current_field['readonly'] = (val not in ('no', 'n'))
                    continue

                m = re.match(r'^Max Length:\s*(.+)', line)
                if m: current_field['max_length'] = m.group(1).strip(); continue

                m = re.match(r'^Pattern:\s*(.+)', line)
                if m: current_field['pattern'] = m.group(1).strip(); continue

                m = re.match(r'^Validation Rule:\s*(.+)', line)
                if m: current_field['validation'] = m.group(1).strip(); continue

                m = re.match(r'^Help Text:\s*(.+)', line)
                if m: current_field['help_text'] = m.group(1).strip(); continue

                m = re.match(r'^Source:\s*(.+)', line)
                if m:
                    current_field['source'] = m.group(1).strip()
                    collecting_options = False
                    continue

                # Cascading
                m = re.match(r'^Cascading:\s*(.+)', line)
                if m:
                    current_field['notes'] += f"Cascading: {m.group(1).strip()}. "
                    continue

                # Condition
                m = re.match(r'^Condition:\s*(.+)', line)
                if m:
                    collecting_condition = True
                    condition_lines = [m.group(1).strip()]
                    continue
                if collecting_condition:
                    if line.startswith('OR ') or line.startswith('AND ') or line.startswith('OR\t'):
                        condition_lines.append(line)
                        continue
                    else:
                        collecting_condition = False

                # Static options
                if re.match(r'^Options\s*\(STATIC', line):
                    collecting_options = True
                    options_lines = []
                    continue
                if collecting_options:
                    m2 = re.match(r'^-\s+label:\s*"(.+?)"\s*\|\s*value:\s*"(.+?)"', line)
                    if m2:
                        options_lines.append(f'{m2.group(1)} ({m2.group(2)})')
                        continue
                    if line.startswith('-') or line.startswith('label:'):
                        continue
                    if not line:
                        collecting_options = False
                        continue
                    # Non-option line: stop collecting
                    collecting_options = False

                # Default
                m = re.match(r'^Default:\s*(.+)', line)
                if m: current_field['notes'] += f"Default: {m.group(1).strip()}. "; continue

            # CONDITIONAL RULES block
            if re.match(r'^CONDITIONAL RULES', line):
                flush_field()
                current_field = None
                continue

            m = re.match(r'^RULE\s+\d+:\s*(.+)', line)
            if m and current_section is not None:
                current_rules.append(m.group(1).strip())
                continue
            # continuation of rule
            if current_rules and line and not re.match(r'^FIELD|^SECTION|^PAGE|^ADD MORE|^NOTE', line):
                if re.match(r'^[A-Z"\']', line) or line.startswith('"') or '=' in line or line.startswith('OR ') or line.startswith('AND '):
                    current_rules[-1] += ' ' + line
                    continue

        flush_section()
        if current_page:
            pages.append(current_page)
        self.pages = pages


# ─── Document Content Builder ─────────────────────────────────────────────────

def build_document(doc: Document, pages: list):
    for page in pages:
        add_page_header(doc, page['num'], page['title'])
        if page.get('note'):
            p = doc.add_paragraph()
            add_run_colored(p, f"ℹ  {page['note'].strip()}", italic=True, font_size=9,
                            color=RGBColor(0x44, 0x55, 0x77))

        for section in page['sections']:
            add_section_header(doc,
                               section['label'],
                               section['title'],
                               section['note'].strip() or None)

            # ADD MORE banner
            if section.get('addmore'):
                ag = section['addmore']
                add_addmore_banner(doc, ag['label'], ag.get('min', 1), ag.get('max', 10))

            if not section['fields']:
                p = doc.add_paragraph()
                add_run_colored(p, "  [Fields to be defined]", italic=True, font_size=9,
                                color=RGBColor(0x88, 0x88, 0x88))
                continue

            tbl = start_field_table(doc)
            for row_idx, fld in enumerate(section['fields']):
                # Build notes column content
                notes_parts = []
                if fld.get('readonly'):
                    notes_parts.append("Read-only")
                if fld.get('max_length'):
                    notes_parts.append(f"Max: {fld['max_length']} chars")
                if fld.get('pattern'):
                    notes_parts.append(f"Pattern: {fld['pattern']}")
                if fld.get('validation'):
                    notes_parts.append(f"Validation: {fld['validation']}")
                if fld.get('source'):
                    notes_parts.append(f"Source: {fld['source']}")
                if fld.get('notes'):
                    notes_parts.append(fld['notes'].strip())
                if fld.get('help_text'):
                    notes_parts.append(fld['help_text'])

                notes_str = "  |  ".join(notes_parts) if notes_parts else ""

                add_field_row(
                    tbl,
                    fld['num'],
                    fld['name'],
                    fld.get('type', 'text'),
                    fld.get('grid', '6'),
                    fld.get('mandatory', 'Yes'),
                    notes_str,
                    row_idx,
                )
                # Options row
                opts = fld.get('options', '') or fld.get('source', '')
                if opts and fld.get('options'):
                    add_options_row(tbl, opts, row_idx)
                # Condition row
                if fld.get('condition'):
                    add_condition_row(tbl, fld['condition'])

            # Conditional rules block
            if section.get('rules'):
                add_rules_block(doc, section['rules'])

        doc.add_paragraph()


# ─── Business Rules Page ──────────────────────────────────────────────────────

BUSINESS_RULES = [
    ("BR-1", "State Level Approval",
     "Investment > ₹15 Crore OR project is in Tumakuru Machine Tool Park\n"
     "→ State Level: MD → Nodal Officer (DD) → Joint Director (JD) → Committee"),
    ("BR-2", "District Level Approval",
     "Investment ≤ ₹15 Crore\n"
     "→ District Level: JD DIC → DD DIC → AD DIC"),
    ("BR-3", "LAC Meeting Required",
     "Land > 10 Acres  OR  (KIADB land source AND land > 2 Acres)\n"
     "→ Application must go to LAC (Land Audit Committee) meeting"),
    ("BR-4", "Committee Level by Investment",
     "₹15 Cr to ₹500 Cr → SLSWCC meeting\n"
     "Investment > ₹500 Cr → HLSWCC meeting"),
    ("BR-5", "VFA Increment for Expansion",
     "For Expansion/Diversification/Modernization:\n"
     "New VFA must be at least 25% more than existing VFA"),
    ("BR-6", "Social Category Benefits",
     "All promoters must belong to SAME social category/gender\n"
     "for social category benefits (SC, ST, Women, Minorities, etc.)"),
    ("BR-7", "DSC Verification",
     "Before CAF submission, system verifies DSC registration.\n"
     "If not registered, investor redirected to DSC registration page."),
    ("BR-8", "No Payment for Offline Transition",
     "No payment required for transitioning offline application to online format."),
    ("BR-9", "Edit Request",
     "Investor can request edit only if application not yet processed by MD.\n"
     "Valid justification must be provided."),
    ("BR-10", "Withdrawal",
     "Investor can withdraw approved CAF by providing proper justification.\n"
     "MD approves/rejects withdrawal request."),
]

APPLICATION_STATUSES = [
    ("Submitted",           "Application submitted by investor"),
    ("Allocated",           "Allocated to Nodal Officer"),
    ("Pending With DD",     "Under review by Nodal Officer (DD)"),
    ("Pending With JD",     "Forwarded to Joint Director"),
    ("Pending With MD",     "Under review by Managing Director"),
    ("Approved For Meeting","MD approved, awaiting committee meeting"),
    ("Meeting Scheduled",   "Committee meeting date fixed"),
    ("Approved",            "Application approved"),
    ("Rejected",            "Application rejected (with reason)"),
    ("Deferred",            "Decision deferred to next meeting"),
    ("Pending With Investor","Query raised, investor must respond"),
    ("Processing Offline",  "MD marked for offline processing"),
    ("Edit Requested",      "Investor requested edit before MD processing"),
    ("Withdrawal Requested","Investor submitted withdrawal request"),
    ("Cancelled",           "MD cancelled an approved CAF"),
]

CAF_ID_FORMAT = "PAN (10) + Sub-District LGD Code (4 digits) + Sequence (2 digits) + Activity (M=Manufacturing / S=Service)"


def add_appendix(doc: Document):
    doc.add_page_break()
    add_heading1(doc, "APPENDIX A — GLOBAL BUSINESS RULES")
    doc.add_paragraph()

    for br_id, br_title, br_desc in BUSINESS_RULES:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'
        c0 = tbl.cell(0, 0)
        c1 = tbl.cell(0, 1)
        c0.width = Inches(1.1)
        c1.width = Inches(5.4)
        set_cell_bg(c0, COLOR_PRIMARY)
        set_cell_bg(c1, RGBColor(0xFF, 0xFD, 0xFC))
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(5)
        p0.paragraph_format.space_after  = Pt(5)
        add_run_colored(p0, f"  {br_id}", bold=True, font_size=10,
                        color=RGBColor(0xFF, 0xFF, 0xFF))
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(5)
        p1.paragraph_format.space_after  = Pt(5)
        add_run_colored(p1, f"{br_title}: ", bold=True, font_size=9.5, color=COLOR_DARK)
        add_run_colored(p1, br_desc, font_size=9)
        doc.add_paragraph()

    doc.add_page_break()
    add_heading1(doc, "APPENDIX B — APPLICATION STATUS FLOW")
    doc.add_paragraph()

    tbl = doc.add_table(rows=1 + len(APPLICATION_STATUSES), cols=2)
    tbl.style = 'Table Grid'
    for i, h in enumerate(["Status", "Description"]):
        cell = tbl.cell(0, i)
        set_cell_bg(cell, COLOR_DARK)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        add_run_colored(p, h, bold=True, font_size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
    for row_i, (status, desc) in enumerate(APPLICATION_STATUSES):
        bg = COLOR_FIELD_ODD if row_i % 2 == 0 else COLOR_FIELD_EVEN
        c0 = tbl.cell(row_i + 1, 0)
        c1 = tbl.cell(row_i + 1, 1)
        c0.width = Inches(1.8)
        c1.width = Inches(4.7)
        set_cell_bg(c0, bg)
        set_cell_bg(c1, bg)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(3)
        p0.paragraph_format.space_after  = Pt(3)
        add_run_colored(p0, status, bold=True, font_size=9, color=COLOR_PRIMARY)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(3)
        p1.paragraph_format.space_after  = Pt(3)
        add_run_colored(p1, desc, font_size=9)

    doc.add_paragraph()
    add_heading2(doc, "CAF ID Format")
    p = doc.add_paragraph()
    add_run_colored(p, "Format:  ", bold=True, font_size=10, color=COLOR_DARK)
    add_run_colored(p, CAF_ID_FORMAT, font_size=10)

    doc.add_page_break()
    add_heading1(doc, "APPENDIX C — DOCUMENT UPLOAD FORMAT REFERENCE (DMS)")
    doc.add_paragraph()
    doc_format_info = [
        ("Field Type",    "file"),
        ("Accepted Formats", ".pdf, .jpg, .jpeg, .png"),
        ("Default Max Size", "5 MB (overridden per field if specified)"),
        ("Mandatory",     "Yes or Conditional (as per Document Master)"),
        ("Storage",       "Uploaded via DMS API, returns document_id stored as field value"),
        ("Grid Span",     "6 (half row per document)"),
        ("Validation Rule Example", '{"accept": ".pdf,.jpg,.jpeg,.png", "maxSizeMB": 5}'),
    ]
    tbl = doc.add_table(rows=len(doc_format_info), cols=2)
    tbl.style = 'Table Grid'
    for i, (key, val) in enumerate(doc_format_info):
        bg = COLOR_FIELD_ODD if i % 2 == 0 else COLOR_FIELD_EVEN
        c0 = tbl.cell(i, 0)
        c1 = tbl.cell(i, 1)
        c0.width = Inches(2.2)
        c1.width = Inches(4.3)
        set_cell_bg(c0, RGBColor(0xF5, 0xF5, 0xF5))
        set_cell_bg(c1, bg)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(3)
        p0.paragraph_format.space_after  = Pt(3)
        add_run_colored(p0, key, bold=True, font_size=9, color=COLOR_DARK)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(3)
        p1.paragraph_format.space_after  = Pt(3)
        add_run_colored(p1, val, font_size=9)

    doc.add_page_break()
    add_heading1(doc, "APPENDIX D — SRS AUTHORING GUIDE (FOR NEW SERVICES)")
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(
        "When onboarding a NEW department service to the FormBuilder system, "
        "follow this exact SRS format. The AI will parse your document and generate "
        "the complete DB-ready JSON automatically — no manual SQL or coding required."
    )
    r.font.size = Pt(10)

    add_heading2(doc, "Template Structure")
    template = """================================================================================
FORMBUILDER SRS — STANDARD FORMAT
================================================================================
FORM METADATA
================================================================================
Form Name       : <Full Form Name>
Department ID   : <integer from m_department table>
Service ID      : <like 1.0 or 591.0 from m_service table>
Form Type ID    : <1=New Application, 2=Renewal, 3=Amendment>
Description     : <Brief description of the form's purpose>
================================================================================

================================================================================
PAGE 1 — <Page Name>
================================================================================

  SECTION 1.1 — <Section Name>
  ──────────────────────────────────────────────────────────────────────────────
  NOTE: <Optional note about this section>

  FIELD 1  | Name: <Field Name>
              Type: text / select / radio / checkbox / date / tel / email /
                    textarea / file / number / multiselect
              Mandatory: Yes / No / Conditional / Auto Populated / Auto Calculated
              Grid: 6 (half row) or 12 (full row)
              Max Length: <number>                     ← for text fields
              Pattern: <regex>                         ← for validated text
              Validation Rule: {"min": 0, "max": 100}  ← for number/date fields
              Default: <default value>                 ← optional
              Is Readonly: Yes / No                    ← for auto-populated
              Source: MASTER (<Master table name>)     ← for MASTER dropdowns
              Options (STATIC):                        ← for STATIC dropdowns
                - label: "Display Text" | value: "snake_case_value"
              Condition: Show only when <Field Name> = <value>
              Help Text: <Guidance text shown below the field>

  CONDITIONAL RULES FOR SECTION 1.1:
  ─────────────────────────────────────
  RULE 1: Show "<Field Name>" when "<Trigger Field>" = "<value>"
  RULE 2: Show "<Field Name>" when "<Trigger Field>" = "val1" OR "val2"
"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(template)
    r.font.name  = "Courier New"
    r.font.size  = Pt(7.5)
    r.font.color.rgb = RGBColor(0x22, 0x22, 0x44)

    add_heading2(doc, "Key Rules to Remember")
    rules_reminder = [
        "Grid 6 = half width (two fields per row). Grid 12 = full width (one field per row).",
        "For address fields, always use Grid: 12.",
        "For repeatable sections (e.g. list of directors), mark as ADD MORE GROUP with min/max rows.",
        "Master dropdowns: write Source: MASTER (<table name>). Do not list options.",
        "Static dropdowns: list all options with label and value pairs.",
        "Values must be lowercase snake_case (e.g. 'new_project', 'public_limited').",
        "Conditional rules: always list at the END of the section (not inline).",
        "Validation Rule must be a JSON object: {\"min\": 0} or {\"accept\": \".pdf\", \"maxSizeMB\": 5}",
        "Auto Populated fields: set Is Readonly: Yes. Auto Calculated: always readonly.",
        "File upload fields: always set Validation Rule with accept and maxSizeMB.",
    ]
    for r_text in rules_reminder:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        add_run_colored(p, r_text, font_size=9.5)


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    srs_txt_path = Path(__file__).parent / "CAF_Full_SRS_v2.txt"
    out_path     = Path(__file__).parent / "CAF_Full_SRS_v2.docx"

    print(f"Reading SRS from: {srs_txt_path}")
    srs_text = srs_txt_path.read_text(encoding="utf-8")

    print("Parsing SRS structure...")
    parser = SRSParser(srs_text)
    print(f"  → Found {len(parser.pages)} pages")
    for pg in parser.pages:
        total_fields = sum(len(s['fields']) for s in pg['sections'])
        print(f"     Page {pg['num']}: {pg['title']} — {len(pg['sections'])} sections, "
              f"{total_fields} fields")

    print("Building DOCX...")
    doc = setup_document()

    add_cover_page(doc)
    add_how_to_use(doc)
    build_document(doc, parser.pages)
    add_appendix(doc)

    doc.save(str(out_path))
    print(f"\n✅  DOCX saved: {out_path}")
    print(f"   Size: {out_path.stat().st_size / 1024:.1f} KB")
    print("\nTo export as PDF: Open in Microsoft Word → File → Save As → PDF")


if __name__ == "__main__":
    main()
